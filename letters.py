from datetime import datetime
from pathlib import Path
from pprint import pprint

import pandas as pd
from docx import Document
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement
from docx.shared import Inches
from peewee import *

from auth_work import oauth
from config import Config
from google_api_calls_abstract import GoogleApiCalls
from utils import Utils


class Letters():

    def __init__(self, db=None, mode=None, work_order_range=None, gsheet_id=None):
        self.main_db = db
        self.spreadsheet_id = gsheet_id 
        if mode == 'testing':
            self.service = test_service
        else:
            self.service = oauth(Config.my_scopes, 'sheet')
        self.work_order_range = work_order_range

    def setup_tables(self, mode=None):
        from backend import PopulateTable
        populate = PopulateTable()
        self.create_tables_list1 = populate.return_tables_list()
        if self.main_db.is_closed() is True:
            self.main_db.connect()
        if mode == 'create_only':
            self.main_db.create_tables(self.create_tables_list1)
        elif mode == 'drop_and_create':
            self.main_db.drop_tables(models=self.create_tables_list1)
            self.main_db.create_tables(self.create_tables_list1)
        return populate

    def fix_name(self, unit, address_bp):
        name = unit[1].split(',')
        name = [n.rstrip().lstrip().capitalize() for n in name]
        unit_wo_prefix = unit[0].split('-')[1]
        return [(' ').join(name[::-1]), address_bp[0], f'Unit #{unit_wo_prefix}', address_bp[1], address_bp[2]]

    def fix_name2(self, name):
        name = name.split(',')
        name = [n.rstrip().lstrip().capitalize() for n in name]
        return (' ').join(name[::-1])
   
    def get_addresses(self):
        from backend import Unit  # import error work around
        self.setup_tables(mode='create_only')
        units = [(row.unit_name, row.tenant, row.last_occupied) for row in Unit.select().namedtuples()]

        if units == []:
            print('Units table is empty')
            exit
        else:
            addresses = []
            for unit in units:
                if unit[0].split('-')[0] == 'CD':
                    addresses.append(self.fix_name(unit, address_bp=Config.ADDRESS_CD))
                elif unit[0].split('-')[0] == 'PT':
                    addresses.append(self.fix_name(unit, address_bp=Config.ADDRESS_PT))
        return addresses
    
    def get_workorders(self, first_dt=None, last_dt=None):
        from backend import WorkOrder
        workorders = [row for row in WorkOrder.select().
                      where(WorkOrder.init_date >= first_dt).
                      where(WorkOrder.init_date <= last_dt).
                      order_by(WorkOrder.init_date).
                      namedtuples()]
        return workorders
    
    def get_workorder_object(self):
        from backend import WorkOrder
        return WorkOrder
            
    def get_doc_title(self, doc, service_docs): #doc is DOCS_FILE_ID
        document = service_docs.documents().get(documentId=doc).execute()

        print('The title of the document is: {}'.format(document.get('title')))

    def create_folder(self, service_drive, folder_name):
        file_metadata = {
        'name': f'{folder_name}',
        'mimeType': 'application/vnd.google-apps.folder'
                        }
        file = service_drive.files().create(body=file_metadata,
                                        fields='id').execute()
        print('Folder ID: %s' % file.get('id'))

    def create_blank_doc(self, service_docs, service_drive, title): 
        # from docs, there is also a method from drive
        title = f'{title}'
        body = {
            'title': title
        }
        doc = service_docs.documents() \
            .create(body=body).execute()
        print('Created document with title: {0}'.format(
            doc.get('title')))
        print('Created document with id: {0}'.format(
            doc.get('id')))

    def insert_text(self, service_docs, doc_id, text1, text2, text3):
        requests = [
            {
                'insertText': {
                    'location': {
                        'index': 1,
                    },
                    'text': text1
                }
            },
                    {
                'insertText': {
                    'location': {
                        'index': 2,
                    },
                    'text': text2
                }
            }
                ]

        result = service_docs.documents().batchUpdate(
            documentId=doc_id, body={'requests': requests}).execute()

    @staticmethod
    def run_script(service, deploy_id, function_name, parameters=None):
        request = {
        'function': function_name, 
        'devMode': True, 
        'parameters': [parameters]
        }

        response = service.scripts().run(
                        body=request,
                        scriptId=deploy_id
                        ).execute()
        pprint(response)

    def get_bal_let_parameters(self, balance_list=None):
        from backend import ProcessingLayer

        player = ProcessingLayer()
        balance_letters = player.show_balance_letter_list_mr_reconciled()
        name_list = []
        unit_list = []
        bal_due_list = []

        for record in balance_letters:
            formatted_name = [name.rstrip().lstrip().capitalize() for name in record.tenant_name.split(',')]
            name_list.append((' ').join(formatted_name[::-1]))
            unit_list.append(record.unit)
            bal_due_list.append(str(record.end_bal))

        formatted_date = datetime.utcnow()
        parameters = {
            'current_date' : datetime.strftime(datetime.utcnow(), '%Y-%m-%d'), 
            'display_month': formatted_date.strftime('%B'),
            'display_year': str(formatted_date.year),
            'unit': unit_list, 
            'name': name_list, 
            'bal_due': bal_due_list, 
            }

        return parameters

    def get_all_archived_work_orders(self):
        print('get_all_archived_work_orders')
        gc = GoogleApiCalls()
        from backend import WorkOrder, db
        from peewee import IntegrityError as PIE
        
        self.setup_tables(mode='create_only')
        
        values = gc.broad_get(service=self.service, spreadsheet_id=self.spreadsheet_id, range=self.work_order_range)
        df = pd.DataFrame(values)
        df.columns = df.iloc[0]
        df = df[1:] # remove first row and set first row as column names
        df = df.rename(columns={df.columns[6]: 'date_completed'})
        df['date'] = pd.to_datetime(df['date'])
        df['date_completed'] = pd.to_datetime(df['date_completed'])
        work_orders_insert_many = []
        work_dict = df.to_dict(orient='records')
        
        work_orders_insert_many = [{
            'name': work_order['name'],
            'init_date': work_order['date'],
            'location': work_order['location'],
            'work_req': work_order['work requested'],
            'notes': work_order['notes'],
            'status': work_order['status'],
            'date_completed': work_order['date_completed'],
            'assigned_to': work_order['assigned to'],
            
        } for work_order in work_dict]
        
        try: 
            # try bulk insert
            query = WorkOrder.insert_many(work_orders_insert_many)
            query.execute()
        except PIE as e:
            print(e, 'BULK INSERT FAILED, attempting atomic inserts')
            
        for item in work_orders_insert_many:
            try:
                with db.atomic():
                    nt = WorkOrder.create(name=item['name'], 
                                          init_date=item['init_date'],
                                          location=item['location'],
                                          work_req=item['work_req'],
                                          notes=item['notes'],
                                          status=item['status'],
                                          date_completed=item['date_completed'],
                                          assigned_to=item['assigned_to'], 
                                        )
            except PIE as e:
                print(e, item)
                nt = WorkOrder.create(name=item['name'], 
                                        init_date=item['init_date'],
                                        location=item['location'],
                                        work_req=item['work_req'],
                                        notes=item['notes'],
                                        status=item['status'],
                                        date_completed=item['date_completed'],
                                        assigned_to='ron/bob/fs', 
                ) 
        

    def bal_let_pprint_parameters(self, parameters):
        print('current date', parameters['current_date'], parameters['display_year'])
        print('display_month', parameters['display_month'])

        bal_data = list(zip(parameters['unit'], parameters['name'], parameters['bal_due']))
        print('\n')
        for item in bal_data:
            row = '{:<9s} {:<20} {:<9s}'.format(item[0], item[1], item[2])
            print(row)
        print('\n')

    def balance_letters(self):
        from backend import ProcessingLayer, db
        '''if testing, I can simulate a prod  database by breakpointing db before teardown'''
        '''if there is an issue, check deployment id'''
        db.connect()
        player = ProcessingLayer()
        parameters = self.get_bal_let_parameters(balance_list=player.show_balance_letter_list_mr_reconciled())
        self.bal_let_pprint_parameters(parameters)
        choice = str(input('Send these results to google script & make balance letters? y/n '))
        if choice == 'y':
            Letters.run_script(service=oauth(Config.my_scopes, 'script'), deploy_id=Config.BALANCE_LETTER_DEPLOY_ID, function_name='balanceLetter', parameters=parameters)
        else:
            print('exiting program')
            exit

    def get_prev_daterange_by_month(self, year, month, prev):
        import calendar
        from datetime import date
        if year is None:
            today = datetime.date.today()
            year = today.year
        if month is None:
            today = datetime.date.today()
            month = today.month
        since = []
        till = []
        for i in range(prev):
            if month == 0:
                year -= 1
                month = 12
            _, num_days = calendar.monthrange(year, month)
            since.append(datetime.strptime(str(date(year, month, 1)), "%Y-%m-%d"))
            till.append(datetime.strptime(str(date(year, month, num_days)), "%Y-%m-%d"))
            month -= 1
        return since, till

    def rent_receipts_plus_balance(self):
        from backend import QueryHC
        print('rent receipts plus balance table')
        query = QueryHC()

        target_date = '2022-07'
        first_dt, last_dt = query.make_first_and_last_dates(date_str=target_date)

        rent_roll = query.get_rent_roll_by_month_at_first_of_month(first_dt=first_dt, last_dt=last_dts)

        '''consider dataframe for merging'''

        breakpoint()






        lookback_months = 6
        today = datetime.today()
        first_date_of_months, last_date_of_months = self.get_prev_daterange_by_month(2022, 7, lookback_months)
        dates = list(zip(first_date_of_months, last_date_of_months))

        
        earliest_date_in_range = first_date_of_months[-1]
        last_date_in_range = last_date_of_months[0]

        '''should do some sort of group by operation, get last date of most recent months & first date of earliest month: should get all payments'''

        total_payments = query.get_payments_by_tenant_by_period(first_dt=earliest_date_in_range, last_dt=last_date_in_range, cumsum=True)

        all_payments = query.        get_payments_by_tenant_by_period(first_dt=earliest_date_in_range, last_dt=last_date_in_range)

        for count, tup in enumerate(all_payments, 1):
            print(count, tup[0])

        breakpoint()
        # set target date

        # get beginning balance for all current tenants
        # get total payments by tenant by period
        # get total charges by tenant by period



        '''
        titles_dict = Utils.get_existing_sheets(oauth(Config.my_scopes, 'sheet'), Config.TEST_RS)
        idx_list = Utils.existing_ids({name:id2 for name, id2 in titles_dict.items() if name != 'intake'})
        sheet_choice = idx_list[int(input('Please select a sheet to make receipts from: '))]
        parameters = {
        'current_date' : datetime.strftime(datetime.utcnow(), '%Y-%m-%d'), 
        'display_month': str(input('Type display month as you wish it to appear? ')),
        'sheet_choice': sheet_choice[1][0], 
        'rent_sheet': Config.TEST_RS, 
        }

        pprint(parameters)
        choice = str(input('Send these results to google script & make receipts? y/n '))
        if choice == 'y':
            Letters.run_script(service=oauth(Config.my_scopes, 'script'), deploy_id=Config.receipts_table_test_deploy_id, function_name="test1", parameters=parameters) 
        else:
            print('exiting program')
            exit
        '''
    def check_rs_status(self):
        titles_dict = Utils.get_existing_sheets(oauth(Config.my_scopes, 'sheet'), Config.TEST_RS)
        idx_list = Utils.existing_ids({name:id2 for name, id2 in titles_dict.items() if name != 'intake'})
        return idx_list

    def rent_receipts_configuration(self):
        idx_list = self.check_rs_status()
        sheet_choice = idx_list[int(input('Please select a sheet to make receipts from: '))]
        parameters = {
        'current_date' : datetime.strftime(datetime.utcnow(), '%Y-%m-%d'), 
        'display_month': str(input('Type display month as you wish it to appear? ')),
        'sheet_choice': sheet_choice[1][0], 
        'rent_sheet': Config.TEST_RS, 
        }
        return parameters

class AddressWriter(Letters):

    testing_save_path = Config.ADDRESS_TESTING_SAVE_PATH

    def __init__(self, db=None, service=None):
        self.main_db = db
        self.header_indent = 4
        self.service = service
        self.df = None

    def show_path(self):
        print(self.testing_save_path)

    def export_to_excel(self, address_list=None):
        self.df = pd.DataFrame(address_list, 
                                columns=['name', 'street', 'unit', 'city', 'zip'], 
                                )
        self.df.set_index('name', inplace=True)
        write_path = self.testing_save_path / Path('testing.xlsx')
        self.df.to_excel(write_path)
        print(f'look for output in {write_path}')

class DocxWriter(Letters):

    default_save_path = Config.TEST_DOCX_BASE
    testing_save_path = Config.PYTEST_DOCX_BASE
    testing_save_name = 'testing.docx'
    workorders_save_name = 'workorders.docx'
    workorders_save_path = Config.WORKORDER_OUTPUT

    def __init__(self, db=None, service=None):
        self.main_db = db
        self.header_indent = 4
        self.service = service

    def load_from_sheet(self, *args, **kwargs):
        gc = GoogleApiCalls()
        return gc.broad_get(service=self.service, spreadsheet_id=args[1], range=args[0] + '!b2:k68')

    def insertHR(self, paragraph):
        p = paragraph._p  # p is the <w:p> XML element
        pPr = p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        pPr.insert_element_before(pBdr,
            'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
            'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
            'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
            'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
            'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
            'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
            'w:pPrChange'
        )
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'auto')
        pBdr.append(bottom)

    def insert_header(self, document=None):
        paragraph = document.add_paragraph('Fall Creek Village I', style='No Spacing')
        paragraph.paragraph_format.left_indent = Inches(self.header_indent)
        paragraph = document.add_paragraph('3515 N. Pennsylvania St.', style='No Spacing')
        paragraph.paragraph_format.left_indent = Inches(self.header_indent)
        paragraph = document.add_paragraph('Indianapolis, IN 46205', style='No Spacing')
        paragraph.paragraph_format.left_indent = Inches(self.header_indent)
        paragraph = document.add_paragraph('(317) 925-5558', style='No Spacing')
        paragraph.paragraph_format.left_indent = Inches(self.header_indent)
        paragraph = document.add_paragraph('TTY: 711 or (800) 743-3333', style='No Spacing')
        paragraph.paragraph_format.left_indent = Inches(self.header_indent)
        paragraph = document.add_paragraph('                                                                                                                                               ')
        self.insertHR(paragraph)

    def format_docx_rent_receipt(self, document=None, parameters=None, r_recs=None):
        for address in r_recs:
            self.insert_header(document)
            paragraph = document.add_paragraph('Date: ' + parameters['current_date'], style='No Spacing')
            paragraph = document.add_paragraph(' ', style='No Spacing')
            paragraph = document.add_paragraph(address[0], style='No Spacing')
            paragraph = document.add_paragraph(address[1], style='No Spacing')
            paragraph = document.add_paragraph(address[2], style='No Spacing')
            city_and_state = address[3] + ' ' + address[4]
            paragraph = document.add_paragraph(city_and_state, style='No Spacing')
            paragraph = document.add_paragraph(' ', style='No Spacing')
            paragraph = document.add_paragraph(f'Dear {address[0]},', style='No Spacing')
            paragraph = document.add_paragraph(' ', style='No Spacing')
            paragraph = document.add_paragraph(' ', style='No Spacing')
            paragraph = document.add_paragraph('Thank you for your rent payment for ' + parameters['display_month'] + ' ' + Config.current_year + '.', style='No Spacing')
            paragraph = document.add_paragraph(' ', style='No Spacing')
            paragraph = document.add_paragraph('Our records show that you paid: $ ' +  address[5] + '.', style='No Spacing')

            paragraph = document.add_paragraph(' ', style='No Spacing')
            paragraph = document.add_paragraph(' ', style='No Spacing')
            paragraph = document.add_paragraph(' ', style='No Spacing')
            paragraph = document.add_paragraph(' ', style='No Spacing')
            paragraph = document.add_paragraph('generated: ' +  parameters['current_date'], style='No Spacing')

            document.add_page_break()
        return document
    
    def format_workorders(self, document=None, parameters=None, records=None):
        for record in records:
            self.insert_header(document)
            paragraph = document.add_paragraph(' ', style='No Spacing')
            paragraph = document.add_paragraph(f'Work Order Date: {record[0]}', style='No Spacing')
            paragraph = document.add_paragraph(' ', style='No Spacing')
            paragraph = document.add_paragraph(f'Tenant Name: {record[1]}', style='No Spacing')
            paragraph = document.add_paragraph(f'Location: {record[2]}', style='No Spacing')
            paragraph = document.add_paragraph(f'Work Requested: {record[3]}', style='No Spacing')
            paragraph = document.add_paragraph(' ', style='No Spacing')
            paragraph = document.add_paragraph(' ', style='No Spacing')
            paragraph = document.add_paragraph(f'Work Status: {record[5]}', style='No Spacing')
            paragraph = document.add_paragraph(' ', style='No Spacing')
            paragraph = document.add_paragraph(f'Notes: {record[4]}', style='No Spacing')
            paragraph = document.add_paragraph(f'Assigned to/Completed by: {record[7]}', style='No Spacing')
            paragraph = document.add_paragraph(' ', style='No Spacing')
            paragraph = document.add_paragraph(' ', style='No Spacing')
            paragraph = document.add_paragraph(' ', style='No Spacing')
            paragraph = document.add_paragraph(' ', style='No Spacing')
            paragraph = document.add_paragraph(' ', style='No Spacing')
            paragraph = document.add_paragraph(' ', style='No Spacing')
            paragraph = document.add_paragraph(' ', style='No Spacing')
            paragraph = document.add_paragraph(f'Date Completed: {record[6]}', style='No Spacing')
            paragraph = document.add_paragraph('Verified by: JW', style='No Spacing')
            paragraph = document.add_picture(Config.image_path, width=Inches(.75), height=Inches(.5))
            paragraph = document.add_paragraph('_______________________________', style='No Spacing')
            paragraph = document.add_paragraph('Generated: ' + parameters['current_date'], style='No Spacing')

            document.add_page_break()
        return document

    def export_workorders_to_docx(self, first_dt=None, last_dt=None):
        print('exporting workorders to docx')
        records = []
        parameters = {'current_date': datetime.now().strftime('%m-%d-%Y')}
        document = Document()
        for workorder in self.get_workorders(first_dt=first_dt, last_dt=last_dt):
            records.append([workorder.init_date, workorder.name, workorder.location, workorder.work_req, workorder.notes, workorder.status, workorder.date_completed, workorder.assigned_to])
           
        document = self.format_workorders(document=document, parameters=parameters, records=records)
        save_name = 'workorders_' + parameters['current_date'] + '.docx'
        save_path = self.workorders_save_path / Path(save_name)
        document.save(save_path)
        return document, save_path 

    def docx_rent_receipts_from_rent_sheet(self, mode=None):
        print('docx rent rent receipts directly from rent sheets')

        if mode == 'testing':
            parameters = {'current_date': '2022-10-27', 'display_month': 'March', 'sheet_choice': '2022-03', 'rent_sheet': '1Z_Qoz-4ehalutipyH2Vj5k-y2b78U69Bc7uXoBKK47Q'}
        else:
            parameters = self.rent_receipts_configuration()

        rs_name_pay = [(self.fix_name2(row[0]), row[9]) for row in self.load_from_sheet(parameters['sheet_choice'], parameters['rent_sheet'])]
        
        self.setup_tables(mode='create_only')
        document = Document()
        r_recs = []
        sum_for_test = []
        for name, pay in rs_name_pay:
            for address in self.get_addresses():
                if name == address[0]:
                    r_recs.append([address[0], address[1], address[2], address[3], address[4], pay])
                    sum_for_test.append(pay)
    
        document = self.format_docx_rent_receipt(document=document, parameters=parameters, r_recs=r_recs)

        if mode == 'testing':
            document.core_properties.title = 'docx_rent_receipts'
            save_path = self.testing_save_path / Path(self.testing_save_name)
        else:
            save_name = 'rent_receipts_' + parameters['current_date'] + '_' + parameters['sheet_choice'] + '.docx'
            save_path = self.default_save_path / Path(save_name)

        document.save(save_path)

        return document, save_path, sum_for_test

   



